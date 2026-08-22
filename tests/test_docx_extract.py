"""Tests for docx structured extraction: sections and heading chains.

Hermetic: docx documents are built in memory with python-docx, so no
external fixtures or network are needed. These tests lock in the
conservative heading detection (built-in Title / Heading N styles only) and
the "merge a section into one segment" behaviour.
"""

from __future__ import annotations

from io import BytesIO

from docx import Document

from lode.ingestion.extract import extract_document, extract_text


def _docx_bytes() -> bytes:
    """A multi-section docx: Title -> body -> H1 -> body -> H2 -> body -> table -> body."""
    doc = Document()
    doc.add_heading("总体报告", level=0)  # Title style (root)
    doc.add_paragraph("前言")
    doc.add_heading("第三章", level=1)
    doc.add_paragraph("第三章正文")
    doc.add_heading("高温耐久性", level=2)
    doc.add_paragraph("详细内容")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "样本"
    table.cell(0, 1).text = "数值"
    table.cell(1, 0).text = "A"
    table.cell(1, 1).text = "1.0"
    doc.add_paragraph("结尾段")
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_extract_document_groups_sections_with_heading_chains() -> None:
    segments = extract_document(_docx_bytes(), ".docx")
    assert segments is not None

    assert len(segments) == 3
    # Section 1: the Title becomes the root heading; its body follows.
    assert segments[0].heading == "总体报告"
    assert segments[0].text == "总体报告\n前言"
    # Section 2: nested under the Title.
    assert segments[1].heading == "总体报告 / 第三章"
    assert segments[1].text == "第三章\n第三章正文"
    # Section 3: deepest level; the table is flattened into the section text.
    assert segments[2].heading == "总体报告 / 第三章 / 高温耐久性"
    assert segments[2].text == "高温耐久性\n详细内容\n样本 | 数值\nA | 1.0\n结尾段"
    # docx has no explicit page numbers yet.
    assert all(segment.page is None for segment in segments)


def test_extract_document_no_headings_single_segment() -> None:
    doc = Document()
    doc.add_paragraph("just body")
    doc.add_paragraph("more body")
    buf = BytesIO()
    doc.save(buf)

    segments = extract_document(buf.getvalue(), ".docx")
    assert segments is not None
    assert len(segments) == 1
    assert segments[0].heading == ""
    assert segments[0].text == "just body\nmore body"


def test_extract_document_skips_empty_paragraphs() -> None:
    doc = Document()
    doc.add_heading("章节", level=1)
    doc.add_paragraph("")
    doc.add_paragraph("内容")
    buf = BytesIO()
    doc.save(buf)

    segments = extract_document(buf.getvalue(), ".docx")
    assert segments is not None
    assert len(segments) == 1
    assert segments[0].heading == "章节"
    assert segments[0].text == "章节\n内容"


def test_extract_text_flattens_docx() -> None:
    text = extract_text(_docx_bytes(), ".docx")
    assert text is not None
    assert "总体报告" in text
    assert "高温耐久性" in text
    assert "样本 | 数值" in text
