"""Shared test doubles and builders for hermetic tests."""

from __future__ import annotations

from io import BytesIO

from docx import Document

from lode.embeddings.base import Embedder
from lode.embeddings.errors import EmbedderUnavailableError
from lode.index.store import FileRecord
from lode.ingestion import Chunk, chunk_digest

DIM = 4


class FakeEmbedder(Embedder):
    """Deterministic in-memory embedder: no network, configurable failure."""

    def __init__(
        self,
        *,
        model_id: str = "test-model",
        dimension: int = DIM,
        fail_model_id: bool = False,
    ) -> None:
        self._model_id = model_id
        self._dimension = dimension
        self.fail_model_id = fail_model_id
        self.dimension_calls = 0

    @property
    def model_id(self) -> str:
        if self.fail_model_id:
            raise RuntimeError("embedding endpoint is down")
        return self._model_id

    @property
    def dimension(self) -> int:
        self.dimension_calls += 1
        return self._dimension

    def embed_documents(self, texts: list[str]) -> list[list[float]]:
        return [[0.1 * (i + 1)] * self._dimension for i in range(len(texts))]

    def embed_query(self, text: str) -> list[float]:
        return [0.1] * self._dimension


class FailingEmbedder(FakeEmbedder):
    """Embedder whose embedding requests always fail (endpoint down).

    Raises the same domain error the real client does, so pipeline tests
    exercise the real failure contract.
    """

    def embed_documents(self, texts: list[str]) -> list[list[float]]:
        raise EmbedderUnavailableError("embedding endpoint is down")

    def embed_query(self, text: str) -> list[float]:
        raise EmbedderUnavailableError("embedding endpoint is down")


def make_chunks(texts: list[str], *, pages: list[int | None] | None = None) -> tuple[list[Chunk], list[list[float]]]:
    chunks = [
        Chunk(digest=chunk_digest(text), text=text, seq=seq, page=pages[seq] if pages else None)
        for seq, text in enumerate(texts)
    ]
    vectors = [[0.1 * (seq + 1), 0.2, 0.3, 0.4] for seq in range(len(texts))]
    return chunks, vectors


def file_record(path: str = "a.txt", *, digest: str = "blake3:aa", size: int = 1) -> FileRecord:
    return FileRecord(path=path, digest=digest, mtime=1.0, size=size)


def make_docx_bytes() -> bytes:
    """A small docx with a Title and a Heading 1 section, as raw bytes."""
    doc = Document()
    doc.add_heading("总体报告", level=0)  # Title (root)
    doc.add_paragraph("前言")
    doc.add_heading("第三章", level=1)
    doc.add_paragraph("高温耐久性内容")
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
